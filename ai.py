"""Claude-powered features + document parsing (handoff §8/§9).

Phase 2 uses:
  - extract_text(): pull raw text from Excel/CSV/PDF/Word/text uploads.
  - normalize_roster(): Claude turns messy roster text into clean athlete rows.
  - google_sheet_csv_url(): turn a share link into a CSV export URL.
Later phases add vision scan-back and the insights chatbot (NB §11: timers get no AI).
"""
import csv
import io
import json
import os

CLAUDE_MODEL = os.environ.get("XC_CLAUDE_MODEL", "claude-sonnet-5")


# ------------------------- Claude client -------------------------
def _client():
    import anthropic  # imported lazily so the app boots without the key
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not set")
    return anthropic.Anthropic(api_key=key)


def claude_chat(system, user, *, max_tokens=4000, model=None, history=None):
    """Chat helper. `history` = optional prior [{role, content}] turns so follow-up
    questions ("what about the girls?") keep their context."""
    msgs = []
    for h in (history or [])[-8:]:
        if isinstance(h, dict) and h.get("role") in ("user", "assistant") and h.get("content"):
            msgs.append({"role": h["role"], "content": str(h["content"])[:4000]})
    msgs.append({"role": "user", "content": user})
    msg = _client().messages.create(
        model=model or CLAUDE_MODEL,
        max_tokens=max_tokens,
        system=system,
        messages=msgs,
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")


def claude_vision(system, prompt, image_bytes, media_type="image/jpeg",
                  *, max_tokens=4000, model=None):
    """Vision helper: send an image + prompt, return the text response."""
    import base64
    b64 = base64.standard_b64encode(image_bytes).decode()
    content = [
        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
        {"type": "text", "text": prompt},
    ]
    msg = _client().messages.create(
        model=model or CLAUDE_MODEL,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": content}],
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")


_HEATSHEET_SYS = (
    "You read a photographed track & field heat sheet. Each row has a lane/section, "
    "a competitor name and/or bib, and a handwritten MARK (a time like 12.34 or "
    "2:05.4 for running events, or a distance/height like 5.42 for field events). "
    "Return ONLY a JSON array, no prose: "
    '[{"bib": <int or null>, "name": "<string or null>", "mark": "<string as written>"}]. '
    "Skip empty/illegible rows. Preserve marks exactly as written."
)


def vision_read_marks(image_bytes, media_type="image/jpeg"):
    """Read handwritten marks off a heat-sheet photo. Returns [{bib,name,mark}]."""
    out = claude_vision(_HEATSHEET_SYS, "Read this heat sheet.", image_bytes,
                        media_type=media_type, max_tokens=4000)
    rows = _find_json_array(out)
    clean = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        mark = r.get("mark")
        if mark in (None, ""):
            continue
        clean.append({"bib": r.get("bib"), "name": r.get("name"), "mark": str(mark)})
    return clean


_SHEET_SYS = (
    "You read a photographed track & field heat sheet. In the TOP-RIGHT corner a sheet "
    "code is printed next to a QR, formatted like 'XCTSHEET E123'. Read that code exactly. "
    "Each row has a lane/section, a competitor name and/or bib, and a handwritten MARK "
    "(a time like 12.34 or 2:05.4 for running, or a distance/height like 5.42 for field). "
    'Return ONLY a JSON object, no prose: {"sheet_code": "<code exactly as printed or null>", '
    '"marks": [{"bib": <int or null>, "name": "<string or null>", "mark": "<string as written>"}]}. '
    "Skip empty/illegible rows. Preserve marks exactly as written."
)


def vision_read_sheet(image_bytes, media_type="image/jpeg"):
    """Read a heat sheet's identifying code AND its marks. Returns
    {"sheet_code": <str|None>, "marks": [{bib,name,mark}]}."""
    out = claude_vision(_SHEET_SYS, "Read this heat sheet's code and marks.", image_bytes,
                        media_type=media_type, max_tokens=4000)
    obj = _find_json_object(out)
    code = obj.get("sheet_code") if isinstance(obj, dict) else None
    marks = []
    for r in (obj.get("marks", []) if isinstance(obj, dict) else []):
        if isinstance(r, dict) and r.get("mark") not in (None, ""):
            marks.append({"bib": r.get("bib"), "name": r.get("name"), "mark": str(r.get("mark"))})
    return {"sheet_code": (str(code).strip() if code else None), "marks": marks}


WAIVER_FIELDS = ("athlete_name", "school", "grade", "district", "date", "dob",
                 "parent_name", "parent_email", "parent_phone",
                 "emergency_name", "emergency_phone")

_WAIVER_TPL_SYS = (
    "You convert an uploaded youth-sports waiver / permission / consent form into a "
    "REUSABLE template. Keep the wording VERBATIM, but replace the specific values that "
    "vary per athlete with placeholders drawn ONLY from this exact list:\n"
    "{{athlete_name}}, {{school}}, {{grade}}, {{district}}, {{date}}, {{dob}}, "
    "{{parent_name}}, {{parent_email}}, {{parent_phone}}, {{emergency_name}}, {{emergency_phone}}\n"
    "Guidance: a student/athlete/child name (or a blank like 'Student Name: ____') -> "
    "{{athlete_name}}; the school or team name -> {{school}}; a grade/year -> {{grade}}; "
    "the district/organization -> {{district}}; a date line -> {{date}}; a date of birth / "
    "DOB -> {{dob}}; the parent or guardian's printed name -> {{parent_name}}; a parent/guardian "
    "email -> {{parent_email}}; a parent/guardian cell/phone -> {{parent_phone}}; an emergency "
    "contact name -> {{emergency_name}}; an emergency contact phone -> {{emergency_phone}}. "
    "Do NOT invent placeholders outside the "
    "list. Leave the actual signature line alone (the app captures the signature "
    "separately). Drop letterhead, logos, and page numbers. Preserve paragraph breaks as "
    r'\n. Return ONLY JSON: {"title": "<short title>", "body": "<template text>"}.'
)


def waiver_template_from_text(raw_text, *, max_chars=20000):
    """Turn a raw uploaded waiver into {title, body} with {{merge fields}}."""
    text = (raw_text or "").strip()
    if not text:
        return {"title": "", "body": ""}
    out = claude_chat(_WAIVER_TPL_SYS, text[:max_chars], max_tokens=4000)
    obj = _find_json_object(out)
    return {"title": str(obj.get("title") or "Team Waiver").strip(),
            "body": str(obj.get("body") or "").strip()}


_FIELD_SYS = (
    "You read a photographed Long Jump / Shot Put sheet. Each row has a BIB and up to THREE "
    "attempt columns (A1, A2, A3). Each attempt is a mark in feet-inches exactly as written "
    "(e.g. 15-06, 5-03, 18-11.5), or 'F' for a foul/scratch, or blank if not taken. "
    'Return ONLY JSON: {"sheet_code": "<XCTSHEET code top-right, or null>", '
    '"rows": [{"bib": <int or null>, "name": "<string or null>", '
    '"attempts": ["<A1>", "<A2>", "<A3>"]}]}. Keep every mark EXACTLY as written — do not '
    "convert, round, or pick a best. Use 'F' for fouls and \"\" for blank. Skip fully empty rows."
)


def vision_read_field(image_bytes, media_type="image/jpeg"):
    """Read a field sheet's code + each athlete's THREE attempts (verbatim strings)."""
    out = claude_vision(_FIELD_SYS, "Read this field-event sheet's code and all attempts.",
                        image_bytes, media_type=media_type, max_tokens=4000)
    obj = _find_json_object(out)
    rows = []
    for r in (obj.get("rows", []) if isinstance(obj, dict) else []):
        if not isinstance(r, dict):
            continue
        atts = [("" if a is None else str(a).strip()) for a in (r.get("attempts") or [])][:3]
        atts += [""] * (3 - len(atts))
        if r.get("bib") is None and not any(atts):
            continue
        rows.append({"bib": r.get("bib"), "name": r.get("name"), "attempts": atts})
    code = obj.get("sheet_code") if isinstance(obj, dict) else None
    return {"sheet_code": (str(code).strip() if code else None), "rows": rows}


_HJ_SYS = (
    "You read a photographed HIGH JUMP results sheet (usually landscape). The HEADER row prints "
    "a series of bar HEIGHTS in feet-inches (e.g. 4-02, 4-04, 4-06 ... up to ~6-02), then a "
    "final 'BEST' column on the far right. Each athlete row has a PRINTED bib and name, then a "
    "handwritten mark in each height column:\n"
    "  O = cleared that height. XO or XXO = missed then cleared (still CLEARED).\n"
    "  X = a miss. XXX (or X then blanks) = three misses, failed out.\n"
    "For EACH athlete, determine BEST = the HIGHEST printed column height whose mark contains an "
    "'O' (i.e. the highest bar they cleared). Read the column heights from the header to know "
    "each column's value. If the far-right BEST box has a height handwritten in it, use THAT "
    "instead of computing. In the TOP-RIGHT corner (may appear rotated) a sheet code prints next "
    "to a QR like 'XCTSHEET E123' — read it exactly. "
    'Return ONLY JSON, no prose: {"sheet_code": "<code exactly as printed, or null>", '
    '"rows": [{"bib": <int or null>, "name": "<string or null>", '
    '"height": "<best CLEARED height in feet-inches (a header value like 5-02), or empty if the '
    'athlete cleared nothing>"}]}. Copy the height exactly as the header prints it. Skip rows '
    "with no bib and no cleared height. Do NOT round or invent heights."
)


def vision_read_hj(image_bytes, media_type="image/jpeg"):
    """Read a High Jump sheet's code + each athlete's best cleared height (+ misses)."""
    out = claude_vision(_HJ_SYS, "Read this high jump sheet's code, bibs, and best heights.",
                        image_bytes, media_type=media_type, max_tokens=4000)
    obj = _find_json_object(out)
    rows = []
    for r in (obj.get("rows", []) if isinstance(obj, dict) else []):
        if not isinstance(r, dict):
            continue
        ht = "" if r.get("height") is None else str(r.get("height")).strip()
        if not ht:
            continue
        m = r.get("misses")
        try:
            m = int(m) if (m is not None and str(m).strip() != "") else None
        except (ValueError, TypeError):
            m = None
        rows.append({"bib": r.get("bib"), "name": r.get("name"), "height": ht, "misses": m})
    code = obj.get("sheet_code") if isinstance(obj, dict) else None
    return {"sheet_code": (str(code).strip() if code else None), "rows": rows}


def _find_json_object(s):
    """Locate and parse the first JSON object from a model response (handles fences)."""
    s = (s or "").strip()
    if s.startswith("```"):
        s = s.split("```", 2)[1]
        if s.startswith("json"):
            s = s[4:]
    a, b = s.find("{"), s.rfind("}")
    if a == -1 or b == -1:
        return {}
    try:
        obj = json.loads(s[a:b + 1])
    except json.JSONDecodeError:
        return {}
    return obj if isinstance(obj, dict) else {}


# ------------------------- document text extraction -------------------------
def extract_text(filename, data):
    """Best-effort raw-text extraction from an uploaded roster file."""
    ext = (filename or "").rsplit(".", 1)[-1].lower()
    if ext in ("csv", "tsv", "txt"):
        return data.decode("utf-8", "replace")
    if ext == "xlsx":
        return _xlsx_text(data)
    if ext == "xls":
        return _xls_text(data)
    if ext == "pdf":
        return _pdf_text(data)
    if ext in ("docx", "doc"):
        return _docx_text(data)
    # Unknown — try utf-8.
    return data.decode("utf-8", "replace")


def _xlsx_text(data):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _xls_text(data):
    import xlrd
    book = xlrd.open_workbook(file_contents=data)
    lines = []
    for sh in book.sheets():
        for r in range(sh.nrows):
            cells = [str(sh.cell_value(r, c)) for c in range(sh.ncols)]
            if any(x.strip() for x in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _pdf_text(data):
    import pdfplumber
    out = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _docx_text(data):
    import docx
    d = docx.Document(io.BytesIO(data))
    lines = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            if any(x.strip() for x in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


# ------------------------- roster normalization -------------------------
_ROSTER_SYS = (
    "You extract a cross-country / track team roster from messy text (spreadsheet "
    "dumps, PDFs, exports). Return ONLY a JSON array, no prose. Each element: "
    '{"name": "First Last", "grade": <int 6-12 or null>, "gender": "M"|"F"|null, '
    '"does_xc": true|false|null, "does_track": true|false|null, '
    '"dob": <date of birth as written, or null>, '
    '"email": <string or null>, "phone": <string or null>, '
    '"parent_name": <string or null>, "parent_email": <string or null>, '
    '"parent_phone": <string or null>, "emergency_name": <string or null>, '
    '"emergency_phone": <string or null>}. '
    "For does_xc/does_track: read columns like 'Cross Country'/'XC' and 'Track' — a Yes, "
    "Y, X, TRUE, ✓, or the sport's name means true; No/blank means false; use null only "
    "when there is no such column at all. "
    "Include the dob/contact/parent/emergency fields ONLY when clearly present in the "
    "source (matching column headers like Date of Birth/DOB, Email, Phone, "
    "Parent/Guardian, Emergency Contact); otherwise use null — never invent them. "
    "Normalize names to 'First Last' with proper capitalization. Infer gender only "
    "if explicit (a column, or M/F/Boys/Girls). Skip header rows, coaches, blanks, "
    "and totals. If a grade is given as 9th/Fr/Freshman etc., map to the integer."
)

# Optional contact fields carried through import when the source has them.
_CONTACT_FIELDS = ("dob", "email", "phone", "parent_name", "parent_email",
                   "parent_phone", "emergency_name", "emergency_phone")


def normalize_roster(raw_text, *, max_chars=20000):
    """Return a list of {name, grade, gender} dicts parsed from raw roster text."""
    text = (raw_text or "").strip()
    if not text:
        return []
    if len(text) > max_chars:
        text = text[:max_chars]
    out = claude_chat(_ROSTER_SYS, text, max_tokens=8000)
    return _parse_json_array(out)


def _find_json_array(s):
    """Locate and parse a JSON array from a model response (handles ``` fences)."""
    s = (s or "").strip()
    if s.startswith("```"):
        s = s.split("```", 2)[1]
        if s.startswith("json"):
            s = s[4:]
    a, b = s.find("["), s.rfind("]")
    if a == -1 or b == -1:
        return []
    try:
        rows = json.loads(s[a:b + 1])
    except json.JSONDecodeError:
        return []
    return rows if isinstance(rows, list) else []


def _parse_json_array(s):
    rows = _find_json_array(s)
    clean = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        name = str(r.get("name", "")).strip()
        if not name:
            continue
        grade = r.get("grade")
        try:
            grade = int(grade) if grade not in (None, "") else None
        except (TypeError, ValueError):
            grade = None
        gender = r.get("gender")
        gender = gender.upper()[0] if isinstance(gender, str) and gender else None
        if gender not in ("M", "F"):
            gender = None
        row = {"name": name, "grade": grade, "gender": gender}
        for k in _CONTACT_FIELDS:
            v = r.get(k)
            row[k] = str(v).strip() if isinstance(v, (str, int)) and str(v).strip() else None
        clean.append(row)
    return clean


# ------------------------- Google Sheet -------------------------
def google_sheet_csv_url(share_url):
    """Turn a Google Sheets share/edit URL into a CSV export URL (or None)."""
    import re
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", share_url or "")
    if not m:
        return None
    sheet_id = m.group(1)
    gid = "0"
    g = re.search(r"[#&?]gid=(\d+)", share_url)
    if g:
        gid = g.group(1)
    return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"


def fetch_google_sheet_text(share_url):
    """Fetch a published/shared Google Sheet as CSV text."""
    import requests
    url = google_sheet_csv_url(share_url)
    if not url:
        raise ValueError("Not a Google Sheets URL")
    r = requests.get(url, timeout=20)
    r.raise_for_status()
    # Normalize to tab-separated for the roster parser.
    rows = list(csv.reader(io.StringIO(r.text)))
    return "\n".join("\t".join(c for c in row) for row in rows if any(row))
